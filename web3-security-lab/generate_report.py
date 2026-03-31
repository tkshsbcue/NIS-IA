#!/usr/bin/env python3
"""Generate compact 15-page content-rich report for Web3 Security Lab"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
style.paragraph_format.space_after = Pt(3)
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.line_spacing = 1.08

for i in range(1, 4):
    hs = doc.styles[f'Heading {i}']
    hs.font.name = 'Calibri'
    hs.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
    hs.font.bold = True
    hs.font.size = Pt([0, 16, 13, 11][i])
    hs.paragraph_format.space_before = Pt([0, 10, 7, 5][i])
    hs.paragraph_format.space_after = Pt(2)


def add_line(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr(); pBdr = OxmlElement('w:pBdr')
    b = OxmlElement('w:bottom')
    b.set(qn('w:val'), 'single'); b.set(qn('w:sz'), '4'); b.set(qn('w:space'), '1'); b.set(qn('w:color'), '999999')
    pBdr.append(b); pPr.append(pBdr)


def add_screenshot(doc, caption):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(0)
    pPr = p._p.get_or_add_pPr(); pBdr = OxmlElement('w:pBdr')
    for name in ['top','left','bottom','right']:
        b = OxmlElement(f'w:{name}')
        b.set(qn('w:val'),'single'); b.set(qn('w:sz'),'4'); b.set(qn('w:space'),'6'); b.set(qn('w:color'),'AAAAAA')
        pBdr.append(b)
    pPr.append(pBdr)
    r = p.add_run('\n[ Insert Screenshot ]\n'); r.font.size = Pt(10); r.font.color.rgb = RGBColor(0x99,0x99,0x99); r.font.italic = True
    cp = doc.add_paragraph(); cp.alignment = WD_ALIGN_PARAGRAPH.CENTER; cp.paragraph_format.space_after = Pt(4)
    r = cp.add_run(caption); r.font.size = Pt(9); r.font.italic = True; r.font.color.rgb = RGBColor(0x66,0x66,0x66)


def set_shading(cell, color):
    s = OxmlElement('w:shd'); s.set(qn('w:fill'), color); s.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(s)


def add_table(doc, headers, rows, col_widths=None):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = h
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(1)
            for r in p.runs: r.font.bold = True; r.font.size = Pt(9); r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        set_shading(c, '1a1a2e')
    for ri, rd in enumerate(rows):
        for ci, ct in enumerate(rd):
            c = t.rows[ri+1].cells[ci]; c.text = str(ct)
            for p in c.paragraphs:
                p.paragraph_format.space_after = Pt(1)
                for r in p.runs: r.font.size = Pt(9)
            set_shading(c, 'F5F5F5' if ri % 2 == 0 else 'FFFFFF')
    if col_widths:
        for row in t.rows:
            for i, w in enumerate(col_widths): row.cells[i].width = Inches(w)
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2)
    return t


def bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(item, style='List Bullet')
        p.paragraph_format.space_after = Pt(1); p.paragraph_format.space_before = Pt(0)
        for r in p.runs: r.font.size = Pt(11)


def code_block(doc, code):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)
    r = p.add_run(code); r.font.name = 'Consolas'; r.font.size = Pt(8); r.font.color.rgb = RGBColor(0x1E,0x29,0x3B)


# ═══════════════════════ TITLE PAGE ═══════════════════════

for _ in range(4): doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Web3 Security Lab'); r.font.size = Pt(32); r.font.bold = True; r.font.color.rgb = RGBColor(0x0F,0x17,0x2A)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Interactive Smart Contract Security Learning Platform'); r.font.size = Pt(14); r.font.color.rgb = RGBColor(0x64,0x74,0x8B)
doc.add_paragraph(); add_line(doc)
for label, val in [('Project Type','Web Application — Blockchain Security Education'),('Tech Stack','React, Vite, TypeScript, Ethers.js, Solidity, TailwindCSS'),('Network','Ethereum Sepolia Testnet'),('Date','March 2026')]:
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f'{label}: '); r.font.bold = True; r.font.size = Pt(10)
    r = p.add_run(val); r.font.size = Pt(10); r.font.color.rgb = RGBColor(0x64,0x74,0x8B)
for _ in range(3): doc.add_paragraph()
for text in ['Prepared By','[Your Name]','[Roll Number / ID]','[Department / Institution]','[Supervisor / Guide Name]']:
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; r = p.add_run(text); r.font.size = Pt(11)

doc.add_page_break()

# ═══════════════════════ 1. ABSTRACT ═══════════════════════

doc.add_heading('1. Abstract', level=1); add_line(doc)
doc.add_paragraph(
    'Smart contracts — self-executing programs deployed on blockchain networks — now govern over $100 billion '
    'in digital assets. Their immutability means vulnerabilities cannot be patched after deployment, making '
    'security education critical. Since 2016, over $6 billion has been lost to exploits including The DAO hack '
    '($60M, reentrancy), Parity Wallet ($150M, access control), and Wormhole Bridge ($320M, signature bypass).'
)
doc.add_paragraph(
    'Web3 Security Lab is an interactive browser-based platform for learning smart contract security through '
    'hands-on exploitation. Users write, compile, deploy, and exploit vulnerable Solidity contracts on the '
    'Ethereum Sepolia testnet using an integrated Monaco editor and MetaMask wallet. Five progressive challenge '
    'levels cover Reentrancy (SWC-107), Access Control (SWC-105), Delegatecall (SWC-112), Integer Overflow '
    '(SWC-101), and Selfdestruct (SWC-132) — each with guided attack steps, hints, and post-exploitation '
    'security education mapped to Module 3.1 (Web Application Vulnerabilities) and Module 3.4 (Secure Code '
    'Review). Built with React, Vite, TypeScript, and TailwindCSS, the modular architecture allows new '
    'levels to be added with zero changes to existing code.'
)
p = doc.add_paragraph()
r = p.add_run('Keywords: '); r.bold = True; r.font.size = Pt(10)
r2 = p.add_run('Smart Contract Security, Ethereum, Solidity, Reentrancy, DeFi, EVM, Blockchain, '
               'Security Scanners, Slither, OWASP ZAP, Sepolia Testnet'); r2.font.size = Pt(10)

# ═══════════════════════ 2. INTRODUCTION ═══════════════════════

doc.add_heading('2. Introduction', level=1); add_line(doc)

doc.add_heading('2.1 Blockchain and Smart Contracts', level=2)
doc.add_paragraph(
    'Blockchain technology, introduced by Satoshi Nakamoto in 2008, created a distributed, tamper-evident '
    'ledger where transactions are grouped into cryptographically linked blocks maintained by a decentralized '
    'network. Bitcoin proved trustless digital value transfer was possible, but its scripting language was '
    'intentionally limited. In 2013, Vitalik Buterin proposed Ethereum — a blockchain with a Turing-complete '
    'programming language enabling smart contracts: self-executing programs that automatically enforce '
    'agreement terms without intermediaries. Ethereum launched in 2015 and now hosts thousands of '
    'decentralized applications (dApps) and DeFi protocols with over $100 billion in total value locked.'
)
doc.add_paragraph(
    'Unlike traditional software where bugs can be patched with server updates, deployed smart contract '
    'code is immutable and publicly visible. Anyone can read the bytecode, find vulnerabilities, and exploit '
    'them irreversibly. This asymmetry — permanent public code with no patching ability — makes smart contract '
    'security uniquely challenging. The financial consequences have been severe:'
)
add_table(doc,
    ['Year', 'Incident', 'Loss', 'Vulnerability'],
    [
        ['2016', 'The DAO', '$60M', 'Reentrancy — recursive withdraw()'],
        ['2017', 'Parity Wallet', '$150M', 'Unprotected delegatecall + kill()'],
        ['2018', 'BEC Token', '$900M mkt cap', 'Integer overflow in batchTransfer()'],
        ['2021', 'Poly Network', '$611M', 'Cross-chain access control bypass'],
        ['2022', 'Wormhole Bridge', '$320M', 'Signature verification bypass'],
        ['2022', 'Ronin Bridge', '$625M', 'Compromised validator keys'],
        ['2023', 'Euler Finance', '$197M', 'Donation attack + liquidation flaw'],
    ],
    [0.5, 1.2, 1.0, 3.3]
)

doc.add_heading('2.2 Problem Statement & Motivation', level=2)
doc.add_paragraph(
    'The security crisis is fundamentally an education crisis — most exploited vulnerabilities are well-documented '
    'with known defenses, yet developers keep shipping vulnerable code. Existing educational resources suffer from: '
    '(1) Theory-practice gap — courses cover concepts but rarely let students execute real exploits on live '
    'blockchains; (2) Fragmented tooling — learning requires juggling Remix, MetaMask, Etherscan, and Hardhat '
    'separately; (3) Steep CTF learning curves with no guided walkthroughs; (4) Missing defensive context — '
    'platforms teach attack but not defense; (5) No academic curriculum mapping. Web3 Security Lab addresses all '
    'five by providing an all-in-one browser platform with integrated editing, deployment, guided exploitation, '
    'and structured post-attack education.'
)

doc.add_heading('2.3 Objectives', level=2)
bullets(doc, [
    'Build an interactive browser platform for learning smart contract security through hands-on exploitation.',
    'Implement in-browser Solidity edit/compile/deploy workflow with MetaMask wallet integration.',
    'Create 5 gamified levels covering Reentrancy, Access Control, Delegatecall, Overflow, Selfdestruct.',
    'Provide step-by-step attack dashboards with state visualization and post-exploit security education.',
    'Demonstrate security scanning tools: Slither, Mythril, OWASP ZAP, Nmap for vulnerability assessment.',
    'Map all content to Module 3.1 (Web App Vulnerabilities) and Module 3.4 (Secure Code Review).',
])

doc.add_page_break()

# ═══════════════════════ 3. THEORETICAL FOUNDATION ═══════════════════════

doc.add_heading('3. Theoretical Foundation', level=1); add_line(doc)

doc.add_heading('3.1 Ethereum Virtual Machine (EVM)', level=2)
doc.add_paragraph(
    'The EVM is a stack-based virtual machine with 256-bit word size that executes smart contract bytecode. '
    'Each contract has persistent storage (256-bit key-value slots), volatile memory, and an execution stack '
    '(max 1024 depth). State variables are assigned to sequential storage slots starting from slot 0 — this '
    'deterministic layout is the root cause of delegatecall storage collisions. Every opcode costs gas '
    '(e.g., ADD=3, SSTORE=5000-20000), metering computation to prevent DoS. When a contract sends ETH '
    'via call{value}(""), execution transfers to the recipient — if the recipient is malicious, it can '
    'call back into the sender before state updates complete, enabling reentrancy.'
)

doc.add_heading('3.2 Solidity Security-Critical Features', level=2)
doc.add_paragraph(
    'Solidity is the primary smart contract language. Security-critical features include: '
    'Visibility modifiers (public/external/internal/private) — forgetting to restrict sensitive functions '
    'is the root cause of access control vulnerabilities. '
    'Fallback and receive() functions — triggered on ETH receipt, attackers place exploit callbacks here. '
    'Delegatecall — executes target code in the caller\'s storage context, enabling proxy upgrades but '
    'causing storage slot collisions when layouts mismatch. '
    'Arithmetic — before Solidity 0.8.0, arithmetic silently overflows/underflows (uint256: 0 - 1 = 2^256 - 1). '
    'Since 0.8.0, automatic overflow checks revert the transaction. '
    'Constructors run once at deployment; for upgradeable proxies, initialize() functions need one-time guards '
    'to prevent attackers from re-initializing and claiming ownership.'
)

# ═══════════════════════ 4. VULNERABILITY DEEP DIVE ═══════════════════════

doc.add_heading('4. Smart Contract Vulnerabilities — Deep Dive', level=1); add_line(doc)

doc.add_heading('4.1 Reentrancy (SWC-107)', level=2)
doc.add_paragraph(
    'The most devastating vulnerability class in blockchain history, responsible for The DAO\'s $60M loss. '
    'A vulnerable withdraw() function sends ETH via call{value: bal}("") before setting balances[msg.sender] = 0. '
    'The attacker\'s contract places a callback to withdraw() inside its receive() function. When ETH arrives, '
    'receive() fires, calling withdraw() again — since the balance hasn\'t been zeroed yet, the check passes '
    'and ETH is sent again, recursively draining the contract. '
    'Defense — Checks-Effects-Interactions (CEI) pattern: check conditions, update state, then make external '
    'calls. Even if re-entered, the balance is already zero. OpenZeppelin\'s ReentrancyGuard adds a mutex lock '
    'that reverts on re-entry.'
)
code_block(doc,
    '// VULNERABLE                              // SECURE (CEI Pattern)\n'
    'function withdraw() public {               function withdraw() public {\n'
    '  uint b = balances[msg.sender];             uint b = balances[msg.sender];\n'
    '  msg.sender.call{value: b}("");  // 1st     balances[msg.sender] = 0;     // Effect FIRST\n'
    '  balances[msg.sender] = 0;       // 2nd     msg.sender.call{value: b}(""); // Interact SECOND\n'
    '}                                          }')

doc.add_heading('4.2 Access Control Failures (SWC-105)', level=2)
doc.add_paragraph(
    'Occurs when privileged functions lack authorization checks. The Parity Wallet hack exploited an '
    'unprotected initWallet() function — anyone could call it to become owner and drain $31M. In a second '
    'incident, an accidental kill() call destroyed the library contract, freezing $150M permanently. '
    'Defense — Set owner in constructor (runs once), use onlyOwner modifiers on all sensitive functions. '
    'For upgradeable proxies, use OpenZeppelin\'s Initializable with an initialized flag preventing re-call.'
)

doc.add_heading('4.3 Delegatecall Storage Collision (SWC-112)', level=2)
doc.add_paragraph(
    'Delegatecall executes target code in the caller\'s storage context. If a Proxy stores owner in slot 0 '
    'and delegatecalls to a Logic contract where slot 0 is uint256 num, calling setNum(attackerAddress) '
    'overwrites the proxy\'s owner. The Parity multisig hack ($150M frozen) exploited this pattern. '
    'Defense — EIP-1967 standardized storage slots: proxy admin data is stored at pseudo-random positions '
    'derived from keccak256 hashes (e.g., bytes32(uint256(keccak256("eip1967.proxy.admin")) - 1)), making '
    'collision with normal variables virtually impossible.'
)

doc.add_heading('4.4 Integer Overflow/Underflow (SWC-101)', level=2)
doc.add_paragraph(
    'Pre-0.8.0 Solidity: uint256 arithmetic wraps silently. The BEC Token exploit used batchTransfer() where '
    'amount * receivers.length overflowed to a small value passing the balance check, while each receiver '
    'got the enormous pre-overflow amount — generating billions of tokens from nothing and crashing $900M '
    'market cap to zero. Defense — use Solidity >= 0.8.0 (built-in checks) or SafeMath library for older versions.'
)

doc.add_heading('4.5 Forced ETH via Selfdestruct (SWC-132)', level=2)
doc.add_paragraph(
    'selfdestruct(address) destroys the calling contract and forcibly sends its ETH to the target — bypassing '
    'receive()/fallback(). Contracts using address(this).balance for logic (e.g., "winner is whoever makes '
    'balance exactly 7 ETH") can be broken by force-sending ETH past the threshold. The game becomes permanently '
    'stuck. Defense — never use address(this).balance for logic; maintain an internal accounting variable '
    '(uint256 totalDeposits) that can only change through explicit deposit/withdraw functions.'
)

doc.add_page_break()

# ═══════════════════════ 5. SECURITY SCANNING TOOLS ═══════════════════════

doc.add_heading('5. Security Scanning Tools — Theory & Execution', level=1); add_line(doc)

doc.add_heading('5.1 Slither — Smart Contract Static Analysis', level=2)
doc.add_paragraph(
    'Slither (Trail of Bits) is the industry-standard static analysis framework for Solidity. It converts '
    'source code into an intermediate representation preserving control/data flow, then runs 90+ detectors '
    'matching vulnerability patterns. It detects reentrancy, unprotected upgrades, unchecked return values, '
    'dangerous delegatecalls, and more — categorized as High/Medium/Low/Informational severity.'
)
code_block(doc,
    'pip3 install slither-analyzer\n'
    'slither contracts/EtherStore.sol                              # Full analysis\n'
    'slither contracts/EtherStore.sol --detect reentrancy-eth      # Specific detector\n'
    'slither contracts/EtherStore.sol --print human-summary        # Summary report')
doc.add_paragraph(
    'On our Level 1 EtherStore contract, Slither reports: Reentrancy in withdraw() — external call before '
    'state update (High), dangerous call value usage (Medium), missing events for state changes (Info).'
)
add_screenshot(doc, 'Fig 5.1: Slither output detecting reentrancy in EtherStore.sol')

doc.add_heading('5.2 Mythril — Symbolic Execution Engine', level=2)
doc.add_paragraph(
    'Mythril (ConsenSys) uses symbolic execution and the Z3 SMT solver to explore execution paths by treating '
    'inputs as symbolic variables. Unlike Slither\'s pattern matching, Mythril can detect vulnerabilities '
    'requiring complex multi-transaction conditions. It analyzes EVM bytecode directly — works even without '
    'source code (e.g., unverified mainnet contracts).'
)
code_block(doc,
    'pip3 install mythril\n'
    'myth analyze contracts/EtherStore.sol --solv 0.8.20           # Source analysis\n'
    'myth analyze --rpc infura-sepolia -a 0xContractAddr           # On-chain bytecode\n'
    'myth analyze contracts/EtherStore.sol --execution-timeout 120  # Extended analysis')
add_screenshot(doc, 'Fig 5.2: Mythril symbolic execution results with SWC-107 finding')

doc.add_heading('5.3 Nmap — Network Scanner', level=2)
doc.add_paragraph(
    'Nmap discovers hosts, open ports, and running services on the infrastructure hosting Web3 frontends '
    'and RPC nodes. A critical finding: exposed Ethereum RPC port 8545 without authentication allows '
    'attackers to call eth_sendTransaction and drain unlocked accounts. Nmap also detects outdated TLS, '
    'unnecessary services, and misconfigured HTTP headers on dApp servers.'
)
code_block(doc,
    'nmap -sV -sC target-server.com                    # Service version + default scripts\n'
    'nmap -p 80,443,8545 -sV target-server.com         # Web + RPC port scan\n'
    'nmap -A -T4 target-server.com                     # Aggressive: OS detect + scripts\n'
    'nmap --script vuln target-server.com               # Vulnerability-specific scripts')
add_screenshot(doc, 'Fig 5.3: Nmap scan showing open ports and service versions')

doc.add_heading('5.4 OWASP ZAP — Web Application Scanner', level=2)
doc.add_paragraph(
    'OWASP ZAP is a DAST (Dynamic Application Security Testing) proxy that intercepts browser-to-server '
    'traffic to detect XSS, CSRF, missing security headers, information leakage, and more. For Web3 dApps, '
    'an XSS vulnerability could inject JavaScript that modifies transaction parameters (changing recipient '
    'address) before MetaMask signing. ZAP\'s Ajax Spider handles React SPAs, and its Active Scanner sends '
    'attack payloads to identify exploitable inputs.'
)
code_block(doc,
    '# GUI: zaproxy → configure browser proxy to localhost:8080\n'
    '# Browse the dApp → ZAP records traffic → Right-click → Active Scan\n'
    'zap-cli quick-scan --self-contained http://localhost:5173     # CLI automated scan\n'
    'zap-cli report -o zap_report.html -f html                    # Generate report')
add_screenshot(doc, 'Fig 5.4: OWASP ZAP alerts categorized by risk level')

doc.add_heading('5.5 Additional Tools', level=2)
doc.add_paragraph(
    'Nikto scans web servers for 6700+ dangerous files, outdated versions, and misconfigurations (exposed .env '
    'files, missing CSP headers). Burp Suite intercepts JSON-RPC calls between the dApp and Ethereum nodes for '
    'manual analysis. MythX combines static analysis, symbolic execution, and fuzzing in a cloud service. '
    'VirusTotal aggregates 70+ antivirus engines to scan wallet software and dApp URLs for malware/phishing.'
)

add_table(doc,
    ['Tool', 'Category', 'Analysis Type', 'Key Strength'],
    [
        ['Slither', 'Smart Contract', 'Static (source)', '90+ detectors, CI-integrated, fast'],
        ['Mythril', 'Smart Contract', 'Symbolic (bytecode)', 'Finds multi-tx complex bugs via Z3'],
        ['Nmap', 'Network', 'Port/service scan', 'Discovers exposed RPC & services'],
        ['OWASP ZAP', 'Web App', 'DAST (dynamic proxy)', 'XSS, CSRF, header analysis for dApps'],
        ['Nikto', 'Web Server', 'Config scanning', 'Detects server misconfigurations'],
        ['Burp Suite', 'Web App', 'Proxy + manual', 'Intercept/replay JSON-RPC calls'],
        ['MythX', 'Smart Contract', 'Cloud multi-engine', 'Combined static + symbolic + fuzzing'],
        ['VirusTotal', 'General', 'Multi-AV scan', '70+ engines, phishing URL detection'],
    ],
    [1.0, 1.0, 1.2, 2.8]
)

doc.add_page_break()

# ═══════════════════════ 6. ARCHITECTURE & IMPLEMENTATION ═══════════════════════

doc.add_heading('6. System Architecture & Implementation', level=1); add_line(doc)

doc.add_heading('6.1 Architecture', level=2)
doc.add_paragraph(
    'Web3 Security Lab is a client-side SPA communicating directly with Ethereum Sepolia via MetaMask. '
    'No backend server — all compilation, state management, and UI run in the browser, minimizing attack '
    'surface. Data flow: user writes Solidity in Monaco Editor → browser compiler generates ABI/bytecode → '
    'ethers.js constructs deployment tx → MetaMask signs → tx broadcast to Sepolia → contract address '
    'returned → user interacts via ethers.js → results displayed in console.'
)
add_screenshot(doc, 'Fig 6.1: System Architecture Diagram')

doc.add_heading('6.2 Tech Stack', level=2)
add_table(doc,
    ['Technology', 'Version', 'Purpose'],
    [
        ['React', '19.2', 'Component UI with hooks, concurrent rendering'],
        ['TypeScript', '5.9', 'Type safety across 29 source files'],
        ['Vite', '8.0', 'Build tool — 10x faster HMR than Webpack'],
        ['TailwindCSS', '4.2', 'Utility-first dark-theme styling'],
        ['ethers.js', '6.16', 'Wallet connection, contract deploy & interaction'],
        ['Monaco Editor', '4.7', 'VS Code engine for Solidity editing'],
        ['Zustand', '5.0', 'State management with localStorage persistence'],
        ['solc-js', '0.8.34', 'Browser-based Solidity compiler'],
    ],
    [1.2, 0.6, 4.2]
)

doc.add_heading('6.3 Key Modules', level=2)
doc.add_paragraph(
    'Code Playground: Monaco Editor with vs-dark theme, Compile/Deploy/Reset toolbar. '
    'Wallet Integration: useWallet hook detects MetaMask, requests accounts, auto-switches to Sepolia '
    '(chainId 0xaa36a7), listens for account/chain changes. '
    'Attack Dashboard: Displays level\'s attack steps as sequential checklist (pending → active → completed) '
    'with state change annotations showing how contract storage evolves during the exploit. '
    'Insights Panel: Post-completion education with vulnerability explanation, secure coding patterns '
    '(Solidity code examples), real-world incident case study, and Module 3.1/3.4 references. '
    'Console: Terminal-style output — cyan/green/red/yellow/purple messages, clickable Etherscan tx links. '
    'Leaderboard: Rankings by levels completed (desc), hints used (asc), localStorage-persisted via Zustand.'
)

doc.add_heading('6.4 Directory Structure', level=2)
code_block(doc,
    'src/\n'
    '├── components/   # 15 components: ui/ layout/ editor/ console/ wallet/ challenge/ leaderboard/\n'
    '├── levels/       # 5 challenge modules with compilable Solidity contracts\n'
    '├── hooks/        # useWallet (MetaMask), useCompiler (compile + deploy)\n'
    '├── lib/          # compiler.ts, deployer.ts, storage.ts\n'
    '├── store/        # Zustand store with persist middleware\n'
    '├── types/        # Level, ConsoleMessage, WalletState, CompilationResult interfaces\n'
    '└── App.tsx       # Header + Sidebar + SplitLayout(Editor|Challenge|Leaderboard | Console|Dashboard|Insights)')

doc.add_page_break()

# ═══════════════════════ 7. CHALLENGE LEVELS ═══════════════════════

doc.add_heading('7. Challenge Levels', level=1); add_line(doc)

add_table(doc,
    ['#', 'Level', 'Difficulty', 'SWC', 'Attack Vector', 'Defense Pattern'],
    [
        ['1', 'Reentrancy', 'Easy', '107', 'Recursive withdraw via receive()', 'CEI + ReentrancyGuard'],
        ['2', 'Ownership Takeover', 'Easy', '105', 'Call unprotected initOwner()', 'Constructor + onlyOwner'],
        ['3', 'Delegatecall', 'Medium', '112', 'Overwrite slot 0 via proxy', 'EIP-1967 storage slots'],
        ['4', 'Integer Overflow', 'Medium', '101', 'Underflow balance in transfer', 'SafeMath / Solidity >=0.8'],
        ['5', 'Selfdestruct', 'Hard', '132', 'Force ETH to break invariant', 'Internal accounting variable'],
    ],
    [0.3, 1.1, 0.7, 0.4, 1.8, 1.7]
)

doc.add_paragraph(
    'Each level contains: compilable vulnerable + attack Solidity contracts, 3 progressive hints (basic → '
    'specific → near-solution), step-by-step attack instructions with state change annotations, and post-exploit '
    'education (vulnerability explanation, secure coding pattern, real-world incident, Module 3.1/3.4 mapping). '
    'The modular architecture means new levels require only a new src/levels/ folder — zero UI/state changes.'
)

# ═══════════════════════ 8. SECURE CODING PATTERNS ═══════════════════════

doc.add_heading('8. Secure Coding Patterns', level=1); add_line(doc)

doc.add_paragraph(
    'This section consolidates the defensive patterns taught across all five levels:'
)
add_table(doc,
    ['Pattern', 'Defends Against', 'How It Works', 'Module'],
    [
        ['Checks-Effects-Interactions', 'Reentrancy', 'Update state before external calls', '3.1 / 3.4'],
        ['ReentrancyGuard', 'Reentrancy', 'Mutex lock reverts on re-entry', '3.1'],
        ['onlyOwner modifier', 'Access Control', 'Constructor sets owner; modifier gates functions', '3.4'],
        ['Initializable', 'Access Control', 'One-time flag prevents re-initialization', '3.4'],
        ['EIP-1967 Storage', 'Delegatecall', 'Hash-derived slot positions prevent collision', '3.4'],
        ['Solidity >= 0.8', 'Integer Overflow', 'Built-in arithmetic revert on overflow', '3.1'],
        ['SafeMath library', 'Integer Overflow', 'Checked add/sub/mul for pre-0.8 contracts', '3.1'],
        ['Internal Accounting', 'Forced ETH', 'Track balance in variable, not address(this).balance', '3.4'],
        ['Pull over Push', 'Multiple', 'Users withdraw; contract never pushes ETH', '3.1 / 3.4'],
    ],
    [1.3, 1.0, 2.3, 0.6]
)

code_block(doc,
    '// Checks-Effects-Interactions — the most important pattern in Solidity:\n'
    'function withdraw() public {\n'
    '    uint bal = balances[msg.sender];        // CHECK\n'
    '    require(bal > 0);\n'
    '    balances[msg.sender] = 0;               // EFFECT (state update BEFORE external call)\n'
    '    (bool ok, ) = msg.sender.call{value: bal}("");  // INTERACTION\n'
    '    require(ok);\n'
    '}')

doc.add_page_break()

# ═══════════════════════ 9. UI/UX & TESTING ═══════════════════════

doc.add_heading('9. UI/UX Design & Testing', level=1); add_line(doc)

doc.add_heading('9.1 Design', level=2)
doc.add_paragraph(
    'Dark theme (#030712 background) for developer comfort. Split-screen layout: left pane (Code Editor / '
    'Challenge / Leaderboard tabs), right pane (Console / Attack Dashboard / Insights tabs). Fixed sidebar '
    '(288px) with level navigation, difficulty badges, completion checkmarks, and progress bar. Monospace '
    'font (JetBrains Mono). Color-coded console: cyan=info, green=success, red=error, yellow=warning, '
    'purple=transactions with clickable Etherscan links.'
)
code_block(doc,
    'App → Header (logo, network badge, wallet) → Sidebar (levels) → SplitLayout\n'
    '  Left Tabs: Editor | Challenge | Leaderboard\n'
    '  Right Tabs: Console | Attack Dashboard | Insights')

doc.add_heading('9.2 Testing Results', level=2)
add_table(doc,
    ['Test', 'Scenario', 'Status'],
    [
        ['Build', 'TypeScript strict-mode: 0 errors / 29 files, Vite build: 565KB JS', 'Pass'],
        ['Wallet', 'MetaMask connect + auto Sepolia switch', 'Pass'],
        ['Editor', 'Load contract, compile, deploy to Sepolia', 'Pass'],
        ['Console', 'Color-coded messages, clickable tx hashes', 'Pass'],
        ['Sidebar', 'Level switching, progress display', 'Pass'],
        ['Hints', 'Progressive 3-tier reveal per level', 'Pass'],
        ['Dashboard', 'Step-through attack with state annotations', 'Pass'],
        ['Insights', 'Post-completion education display', 'Pass'],
        ['Leaderboard', 'Rankings update after level completion', 'Pass'],
    ],
    [0.8, 3.2, 0.6]
)

doc.add_page_break()

# ═══════════════════════ 10. SCREENSHOTS ═══════════════════════

doc.add_heading('10. Screenshots', level=1); add_line(doc)

for cap in [
    'Fig 10.1: Main Interface — sidebar, code editor, console split-screen',
    'Fig 10.2: MetaMask wallet connected with Sepolia network badge',
    'Fig 10.3: Code Editor with vulnerable EtherStore contract',
    'Fig 10.4: Console showing compilation result and deployed contract address',
    'Fig 10.5: Challenge Panel with hints, attack steps, and completion button',
    'Fig 10.6: Attack Dashboard — step-by-step reentrancy exploit',
    'Fig 10.7: Security Insights — vulnerability explanation, CEI pattern, DAO hack reference',
    'Fig 10.8: Leaderboard rankings',
]:
    add_screenshot(doc, cap)

doc.add_page_break()

# ═══════════════════════ 11. FUTURE + CONCLUSION ═══════════════════════

doc.add_heading('11. Future Enhancements', level=1); add_line(doc)
bullets(doc, [
    'New levels: flash loan attacks, front-running/MEV, oracle manipulation, signature replay, proxy hijacking.',
    'Full solc-js Web Worker compilation with version selection and optimization settings.',
    'Attack replay animation with visual state transition diagrams.',
    'Firebase backend for cloud-synced accounts, global leaderboard, instructor analytics.',
    'Multi-chain support: Polygon Amoy, Arbitrum Sepolia, Base Sepolia.',
    'On-chain ERC-721 completion certificates as verifiable credentials.',
    'Slither/Mythril integration: run automated scans on user code directly from the browser.',
])

doc.add_heading('12. Conclusion', level=1); add_line(doc)
doc.add_paragraph(
    'The smart contract security crisis is an education crisis — over $6 billion lost to well-known, preventable '
    'vulnerability classes. Web3 Security Lab addresses this by providing an interactive platform where learners '
    'exploit vulnerable contracts on Sepolia, then immediately study the defensive patterns that prevent each '
    'attack. The "learn by breaking" pedagogy mirrors professional auditing methodology and produces deeper '
    'understanding than passive instruction.'
)
doc.add_paragraph(
    'Five challenge levels cover the most impactful vulnerability classes (Reentrancy, Access Control, '
    'Delegatecall, Integer Overflow, Selfdestruct) with compilable Solidity contracts, guided step-by-step '
    'attacks, progressive hints, and comprehensive security education. The security tool demonstrations '
    '(Slither, Mythril, OWASP ZAP, Nmap) extend learning to automated vulnerability assessment. Built on '
    'modular React/Vite/TypeScript architecture with 29 source files, new levels require only a single new '
    'module. Direct mapping to Module 3.1 and Module 3.4 makes the platform suitable for academic integration. '
    'In an ecosystem where one line of insecure code causes millions in irreversible losses, hands-on security '
    'education is not optional — it is an imperative.'
)

# ═══════════════════════ REFERENCES ═══════════════════════

doc.add_heading('13. References', level=1); add_line(doc)
refs = [
    'Atzei, N., Bartoletti, M., & Cimoli, T. (2017). "A Survey of Attacks on Ethereum Smart Contracts." Intl. Conf. Principles of Security and Trust, pp. 164-186.',
    'Nakamoto, S. (2008). "Bitcoin: A Peer-to-Peer Electronic Cash System." bitcoin.org/bitcoin.pdf',
    'Buterin, V. (2013). "Ethereum: A Next-Generation Smart Contract and Decentralized Application Platform."',
    'Praitheeshan, P. et al. (2019). "Security Analysis Methods on Ethereum Smart Contract Vulnerabilities." arXiv:1908.08605.',
    'SWC Registry. (2024). "Smart Contract Weakness Classification." swcregistry.io',
    'OpenZeppelin. (2023). "Ethernaut." ethernaut.openzeppelin.com',
    'Trail of Bits. (2024). "Slither: Solidity Static Analysis." github.com/crytic/slither',
    'ConsenSys. (2024). "Mythril: EVM Security Analysis." github.com/ConsenSys/mythril',
    'OWASP. (2024). "ZAP — Zed Attack Proxy." zaproxy.org',
    'Nmap Project. (2024). "Nmap Network Scanner." nmap.org',
    'Ethereum Foundation. (2024). "ethers.js v6." docs.ethers.org/v6',
    'EIP-1967: Proxy Storage Slots. (2019). eips.ethereum.org/EIPS/eip-1967',
    'ConsenSys. (2024). "Smart Contract Best Practices." consensys.github.io/smart-contract-best-practices',
    'Solidity Documentation. (2024). "Security Considerations." docs.soliditylang.org',
]
for i, ref in enumerate(refs, 1):
    p = doc.add_paragraph()
    r = p.add_run(f'[{i}] '); r.bold = True; r.font.size = Pt(9)
    r = p.add_run(ref); r.font.size = Pt(9)

# ═══════════════════════ SAVE ═══════════════════════

path = '/Users/kumartanay/NIS-IA/web3-security-lab/Web3_Security_Lab_Project_Report.docx'
doc.save(path)
print(f'Report saved: {path}')
